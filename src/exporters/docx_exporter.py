"""
Word 문서 내보내기
"""

from pathlib import Path
from typing import Optional
import io

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    DocxDocument = None

try:
    from PIL import Image
    import numpy as np
except ImportError:
    Image = None

from .base import BaseExporter
from ..core.builder import DocumentContent, ContentBlock, PageContent


class DocxExporter(BaseExporter):
    """Word 문서로 내보내기"""

    def __init__(
        self,
        page_width: float = 21.0,  # cm
        page_height: float = 29.7,  # cm (A4)
        margin: float = 2.54,  # cm
        font_name: str = "맑은 고딕",
        font_size: int = 11
    ):
        if DocxDocument is None:
            raise ImportError("python-docx가 필요합니다: pip install python-docx")

        self.page_width = page_width
        self.page_height = page_height
        self.margin = margin
        self.font_name = font_name
        self.font_size = font_size

    @property
    def format_name(self) -> str:
        return "Microsoft Word"

    @property
    def file_extension(self) -> str:
        return ".docx"

    def export(self, content: DocumentContent, output_path: Path) -> Path:
        """DocumentContent를 Word 문서로 내보내기"""
        output_path = Path(output_path)
        if not output_path.suffix:
            output_path = output_path.with_suffix(self.file_extension)

        doc = DocxDocument()

        # 페이지 설정
        self._setup_page(doc)

        # 각 페이지 처리
        for page_content in content.pages:
            self._add_page_content(doc, page_content)

            # 페이지 구분 (마지막 페이지 제외)
            if page_content != content.pages[-1]:
                doc.add_page_break()

        # 저장
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return output_path

    def _setup_page(self, doc: DocxDocument) -> None:
        """페이지 설정"""
        section = doc.sections[0]
        section.page_width = Cm(self.page_width)
        section.page_height = Cm(self.page_height)
        section.left_margin = Cm(self.margin)
        section.right_margin = Cm(self.margin)
        section.top_margin = Cm(self.margin)
        section.bottom_margin = Cm(self.margin)

    def _add_page_content(self, doc: DocxDocument, page: PageContent) -> None:
        """페이지 콘텐츠 추가"""
        for block in page.blocks:
            self._add_block(doc, block)

    def _add_block(self, doc: DocxDocument, block: ContentBlock) -> None:
        """콘텐츠 블록 추가"""
        if block.block_type == "text":
            self._add_text_block(doc, block)
        elif block.block_type == "math":
            self._add_math_block(doc, block)
        elif block.block_type == "image":
            self._add_image_block(doc, block)
        elif block.block_type == "table":
            self._add_table_block(doc, block)

    def _add_text_block(self, doc: DocxDocument, block: ContentBlock) -> None:
        """텍스트 블록 추가"""
        if not block.content:
            return

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(str(block.content))
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)

        # 손글씨인 경우 이탤릭 표시 (선택적)
        if block.metadata.get("is_handwritten"):
            run.italic = True

    def _add_math_block(self, doc: DocxDocument, block: ContentBlock) -> None:
        """수식 블록 추가"""
        latex = block.metadata.get("latex") or block.content

        if not latex:
            return

        # LaTeX 수식을 이미지로 변환하여 삽입
        # 또는 텍스트로 삽입
        paragraph = doc.add_paragraph()

        # 수식을 코드 스타일로 표시
        run = paragraph.add_run(f"[수식] {latex}")
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _add_image_block(self, doc: DocxDocument, block: ContentBlock) -> None:
        """이미지 블록 추가"""
        if block.content is None:
            return

        if Image is None:
            return

        try:
            # numpy array를 이미지로 변환
            if isinstance(block.content, np.ndarray):
                pil_image = Image.fromarray(block.content)
            else:
                return

            # 메모리에 저장
            img_buffer = io.BytesIO()
            pil_image.save(img_buffer, format="PNG")
            img_buffer.seek(0)

            # 이미지 크기 계산
            max_width = self.page_width - (2 * self.margin)
            img_width_cm = min(pil_image.width / 96 * 2.54, max_width)

            # 문서에 추가
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(img_buffer, width=Cm(img_width_cm))

        except Exception as e:
            # 이미지 추가 실패 시 플레이스홀더 텍스트
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("[이미지]")
            run.italic = True

    def _add_table_block(self, doc: DocxDocument, block: ContentBlock) -> None:
        """표 블록 추가"""
        table_data = block.content

        if not table_data or not isinstance(table_data, list):
            # 표 데이터가 없으면 텍스트로 표시
            if block.metadata.get("text"):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(str(block.metadata.get("text")))
            return

        if not table_data[0]:
            return

        # 표 생성
        rows = len(table_data)
        cols = len(table_data[0])

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        # 데이터 채우기
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_value in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = str(cell_value) if cell_value else ""

                    # 폰트 설정
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.font_name
                            run.font.size = Pt(self.font_size - 1)
